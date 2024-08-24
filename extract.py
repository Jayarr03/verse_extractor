import tkinter as tk
from tkinter import filedialog, messagebox
from docx import Document
from docx.shared import Pt, RGBColor


def extract_full_rows(docx_path, output_file, output_type):
    # Load the document
    doc = Document(docx_path)

    extracted_lines = []

    # Check tables for rows with a cell starting with a number followed by '#h'
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.lstrip().startswith(tuple(f"{i}#h" for i in range(1, 101))):
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell)
                    extracted_lines.append(row_data)
                    break

    # Save the extracted lines
    if output_type == 'txt':
        with open(output_file, 'w') as f:
            for row_data in extracted_lines:
                line = ' '.join([cell.text.strip() for cell in row_data])
                f.write(line + '\n')
        messagebox.showinfo("Success", f"Data saved to {output_file}")

    elif output_type == 'docx':
        new_doc = Document()
        for row_data in extracted_lines:
            para = new_doc.add_paragraph()
            for cell in row_data:
                for para in cell.paragraphs:
                    for run in para.runs:
                        new_run = para.add_run(run.text)
                        new_run.bold = run.bold
                        new_run.italic = run.italic
                        new_run.underline = run.underline
                        new_run.font.size = run.font.size
                        if run.font.color.rgb:
                            new_run.font.color.rgb = run.font.color.rgb
                        if run.font.name:
                            new_run.font.name = run.font.name
            para.add_run(' ')
        new_doc.save(output_file)
        messagebox.showinfo("Success", f"Data saved to {output_file}")


def select_file():
    file_path = filedialog.askopenfilename(filetypes=[("Word Documents", "*.docx")])
    file_entry.delete(0, tk.END)
    file_entry.insert(0, file_path)


def run_extraction():
    file_path = file_entry.get()
    output_type = output_var.get()

    if not file_path:
        messagebox.showwarning("Input Error", "Please select a file.")
        return

    output_file = filedialog.asksaveasfilename(defaultextension=f".{output_type}",
                                               filetypes=[("Text Files", "*.txt"), ("Word Documents", "*.docx")])

    if output_file:
        extract_full_rows(file_path, output_file, output_type)


# Create the main window
root = tk.Tk()
root.title("Text Extractor")

# File selection
file_label = tk.Label(root, text="Select File:")
file_label.grid(row=0, column=0, padx=10, pady=10)

file_entry = tk.Entry(root, width=50)
file_entry.grid(row=0, column=1, padx=10, pady=10)

file_button = tk.Button(root, text="Browse...", command=select_file)
file_button.grid(row=0, column=2, padx=10, pady=10)

# Output type selection
output_var = tk.StringVar(value="txt")

txt_radio = tk.Radiobutton(root, text="Text File (.txt)", variable=output_var, value="txt")
txt_radio.grid(row=1, column=0, padx=10, pady=10, sticky="w")

docx_radio = tk.Radiobutton(root, text="Word Document (.docx)", variable=output_var, value="docx")
docx_radio.grid(row=1, column=1, padx=10, pady=10, sticky="w")

# Run button
run_button = tk.Button(root, text="Run Extraction", command=run_extraction)
run_button.grid(row=2, column=0, columnspan=3, padx=10, pady=20)

# Run the main loop
root.mainloop()
